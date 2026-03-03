from fastapi import APIRouter, Header, HTTPException
from fastapi.responses import StreamingResponse
from supabase import create_client
from app.config import SUPABASE_URL, SUPABASE_SERVICE_KEY
from docx import Document
from docx.shared import Pt, RGBColor
import io

router = APIRouter()

def get_supabase(token: str):
    return create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)


def get_run_and_questions(supabase, run_id: str):
    run_result = supabase.table("questionnaire_runs").select("*").eq("id", run_id).execute()
    if not run_result.data:
        raise HTTPException(status_code=404, detail="Run not found")
    run = run_result.data[0]

    questions_result = supabase.table("questions").select("*").eq("run_id", run_id).order("question_number").execute()
    questions = questions_result.data
    return run, questions


def wrap_text(text: str, max_chars: int = 95) -> list[str]:
    words = (text or "").split()
    if not words:
        return [""]

    lines = []
    current = []
    current_len = 0
    for word in words:
        new_len = current_len + len(word) + (1 if current else 0)
        if new_len > max_chars:
            lines.append(" ".join(current))
            current = [word]
            current_len = len(word)
        else:
            current.append(word)
            current_len = new_len

    if current:
        lines.append(" ".join(current))

    return lines


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_simple_pdf(lines: list[str]) -> bytes:
    page_width = 612
    page_height = 792
    margin_x = 50
    margin_y = 50
    line_height = 14

    pages: list[list[str]] = []
    current_page: list[str] = []
    max_lines_per_page = (page_height - 2 * margin_y) // line_height

    for line in lines:
        if len(current_page) >= max_lines_per_page:
            pages.append(current_page)
            current_page = []
        current_page.append(line)

    if current_page or not pages:
        pages.append(current_page)

    objects: list[bytes] = []

    def add_object(content: str) -> int:
        obj_num = len(objects) + 1
        obj = f"{obj_num} 0 obj\n{content}\nendobj\n".encode("latin-1", errors="replace")
        objects.append(obj)
        return obj_num

    font_obj = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    page_obj_ids = []
    content_obj_ids = []

    for page_lines in pages:
        commands = ["BT", "/F1 11 Tf"]
        y = page_height - margin_y
        for line in page_lines:
            safe = _pdf_escape(line)
            commands.append(f"1 0 0 1 {margin_x} {y} Tm ({safe}) Tj")
            y -= line_height
        commands.append("ET")
        stream_data = "\n".join(commands).encode("latin-1", errors="replace")
        content_obj = add_object(
            f"<< /Length {len(stream_data)} >>\nstream\n{stream_data.decode('latin-1')}\nendstream"
        )
        content_obj_ids.append(content_obj)

        page_obj = add_object(
            "<< /Type /Page /Parent {PAGES_REF} 0 R "
            f"/MediaBox [0 0 {page_width} {page_height}] "
            f"/Resources << /Font << /F1 {font_obj} 0 R >> >> "
            f"/Contents {content_obj} 0 R >>"
        )
        page_obj_ids.append(page_obj)

    kids = " ".join(f"{pid} 0 R" for pid in page_obj_ids)
    pages_obj_num = add_object(f"<< /Type /Pages /Count {len(page_obj_ids)} /Kids [ {kids} ] >>")

    for idx, obj in enumerate(objects):
        if b"{PAGES_REF}" in obj:
            objects[idx] = obj.replace(b"{PAGES_REF}", str(pages_obj_num).encode("ascii"))

    catalog_obj = add_object(f"<< /Type /Catalog /Pages {pages_obj_num} 0 R >>")

    output = io.BytesIO()
    output.write(b"%PDF-1.4\n")

    offsets = [0]
    for obj in objects:
        offsets.append(output.tell())
        output.write(obj)

    xref_start = output.tell()
    output.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.write(b"0000000000 65535 f \n")
    for i in range(1, len(objects) + 1):
        output.write(f"{offsets[i]:010d} 00000 n \n".encode("ascii"))

    output.write(
        (
            "trailer\n"
            f"<< /Size {len(objects) + 1} /Root {catalog_obj} 0 R >>\n"
            f"startxref\n{xref_start}\n%%EOF"
        ).encode("ascii")
    )

    return output.getvalue()

@router.get("/{run_id}")
async def export_run(run_id: str, authorization: str = Header(...)):
    token = authorization.replace("Bearer ", "")
    supabase = get_supabase(token)

    run, questions = get_run_and_questions(supabase, run_id)

    # Create Word document
    doc = Document()

    # Title
    title = doc.add_heading(f"EduVault - {run['title']}", 0)
    title.alignment = 1

    # Coverage Summary
    doc.add_heading("Coverage Summary", level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"Total Questions: {run['coverage_total']}\n").bold = True
    summary.add_run(f"Answered with Citations: {run['coverage_answered']}\n").bold = True
    summary.add_run(f"Not Found in References: {run['coverage_not_found']}\n").bold = True

    doc.add_paragraph("")

    # Questions and Answers
    doc.add_heading("Questions & Answers", level=1)

    for q in questions:
        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(q['question_text'] or "")
        q_run.bold = True
        q_run.font.size = Pt(12)

        # Answer
        a_para = doc.add_paragraph()
        a_para.add_run("Answer: ").bold = True
        a_para.add_run(q['answer'] or "Not found in references.")

        # Confidence
        confidence = q.get('confidence', 0)
        c_para = doc.add_paragraph()
        c_run = c_para.add_run(f"Confidence: {int(confidence * 100)}%")
        c_run.italic = True
        c_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Citations
        citations = q.get('citations', [])
        if citations and citations[0]:
            cit_para = doc.add_paragraph()
            cit_para.add_run("Citations: ").bold = True
            cit_para.add_run(", ".join(citations))

        # Edited badge
        if q.get('is_edited'):
            e_para = doc.add_paragraph()
            e_run = e_para.add_run("✏️ Manually edited")
            e_run.italic = True

        doc.add_paragraph("─" * 60)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"eduvault_{run['title'].replace(' ', '_')}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/{run_id}/pdf")
async def export_run_pdf(run_id: str, authorization: str = Header(...)):
    token = authorization.replace("Bearer ", "")
    supabase = get_supabase(token)

    run, questions = get_run_and_questions(supabase, run_id)

    lines = [
        f"EduVault - {run['title']}",
        "",
        "Coverage Summary",
        f"Total Questions: {run['coverage_total']}",
        f"Answered with Citations: {run['coverage_answered']}",
        f"Not Found in References: {run['coverage_not_found']}",
        "",
        "Questions & Answers",
        "",
    ]

    for q in questions:
        lines.extend(wrap_text(q.get("question_text") or ""))
        lines.extend(wrap_text(f"Answer: {q.get('answer') or 'Not found in references.'}"))
        lines.append(f"Confidence: {int((q.get('confidence') or 0) * 100)}%")

        citations = q.get("citations") or []
        citations_text = ", ".join(citations) if citations and citations[0] else "None"
        lines.extend(wrap_text(f"Citations: {citations_text}"))

        if q.get("is_edited"):
            lines.append("Manually edited")

        lines.append("-" * 80)
        lines.append("")

    pdf_bytes = build_simple_pdf(lines)
    buffer = io.BytesIO(pdf_bytes)

    filename = f"eduvault_{run['title'].replace(' ', '_')}.pdf"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
